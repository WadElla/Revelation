import os
import time
import psutil
import hashlib
import pandas as pd
from transformers import AutoTokenizer
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from nltk.translate.meteor_score import meteor_score
from rouge_score import rouge_scorer
from bert_score import score as bert_score
from pynvml import (
    nvmlInit, nvmlShutdown, nvmlDeviceGetHandleByIndex,
    nvmlDeviceGetMemoryInfo, nvmlDeviceGetUtilizationRates,
    nvmlDeviceGetComputeRunningProcesses
)
from langchain.schema import Document
from langchain_community.vectorstores import Chroma
from langchain_community.llms import Ollama
from get_embedding_function import get_embedding_function
from docx import Document as DocxDocument

from langchain_google_genai import ChatGoogleGenerativeAI
import os
os.environ["GOOGLE_API_KEY"] = "AIzaSyCtowOePpR13zVQ9n2IaKgJSXMBMMNFgG0"

import google.generativeai as genai

genai.configure(api_key="AIzaSyCtowOePpR13zVQ9n2IaKgJSXMBMMNFgG0")

# Config
CHROMA_PATH = None
K_DENSE = 30
TOP_K_FINAL = 7
MODEL_NAME = "llama3.2:3b"
DOCX_QA_PATH = "Evaluation_results/Backdoor_test3.docx"

# ----------------- Monitoring utils -----------------
def initialize_nvml(): nvmlInit()
def shutdown_nvml(): nvmlShutdown()

def get_process_gpu_usage(gpu_index=0):
    pid = os.getpid()
    handle = nvmlDeviceGetHandleByIndex(gpu_index)
    for p in nvmlDeviceGetComputeRunningProcesses(handle):
        if p.pid == pid:
            return {"gpu_memory_used": p.usedGpuMemory / (1024 ** 2), "gpu_utilization": None}
    return {"gpu_memory_used": 0.0, "gpu_utilization": None}

def get_device_gpu_usage(gpu_index=0):
    handle = nvmlDeviceGetHandleByIndex(gpu_index)
    mem = nvmlDeviceGetMemoryInfo(handle)
    util = nvmlDeviceGetUtilizationRates(handle)
    return {"gpu_memory_used": mem.used / (1024 ** 2), "gpu_utilization": util.gpu}

# ----------------- Answer evaluation -----------------
def evaluate_answer(pred, ref):
    ref_toks, pred_toks = ref.split(), pred.split()
    P, R, F1 = bert_score([pred], [ref], lang="en", verbose=False)
    rouge = rouge_scorer.RougeScorer(['rouge1','rouge2','rougeL'], use_stemmer=True).score(ref, pred)
    return {
        "bert_Precision": P.mean().item(), "bert_Recall": R.mean().item(), "bert_F1": F1.mean().item(),
        "rouge1": rouge['rouge1'].fmeasure, "rouge2": rouge['rouge2'].fmeasure, "rougeL": rouge['rougeL'].fmeasure,
        "bleu": sentence_bleu([ref_toks], pred_toks, smoothing_function=SmoothingFunction().method1),
        "meteor": meteor_score([ref_toks], pred_toks)
    }

def calculate_response_size(text):
    try:
        tokens = AutoTokenizer.from_pretrained("gpt2").tokenize(text)
        return sum(len(t.encode()) for t in tokens), len(tokens)
    except:
        return 0, 0

def load_questions(path):
    doc = DocxDocument(path)
    return {p.text.split(":",1)[0].strip(): p.text.split(":",1)[1].strip()
            for p in doc.paragraphs if ":" in p.text}

# ----------------- Main -----------------
def main():
    initialize_nvml()

    global CHROMA_PATH
    if not CHROMA_PATH:
        with open("chroma_sessions/latest_session.txt") as f:
            CHROMA_PATH = f.read().strip()

    embed_fn = get_embedding_function()
    db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embed_fn)
    raw = db.get(include=["documents", "metadatas"])
    all_docs = [Document(page_content=t, metadata=m) for t, m in zip(raw["documents"], raw["metadatas"])]

    #model = ChatGoogleGenerativeAI(model="gemini-2.5-pro", temperature=0.6)
    model = Ollama(model=MODEL_NAME, temperature=0.6)

    questions = load_questions(DOCX_QA_PATH)

    result_rows, metric_rows = [], []
    process = psutil.Process()

    for question, answer in questions.items():
        mem_before = process.memory_info().rss / (1024**2)
        proc_gpu_before = get_process_gpu_usage()
        device_gpu_before = get_device_gpu_usage()
        cpu_before = process.cpu_times()
        t0 = time.time()

        # Dense-only retrieval (fixed scoring/ordering)
        try:
            # Prefer relevance scores: 0..1, higher is better
            results_db = db.similarity_search_with_relevance_scores(question, k=K_DENSE)  # [(doc, score)]
            scored = [(float(score), doc) for doc, score in results_db]
        except Exception:
            # Fallback to distances: convert distance -> similarity
            results_db = db.similarity_search_with_score(question, k=K_DENSE)  # [(doc, distance)]
            scored = []
            for doc, dist in results_db:
                try:
                    dist = float(dist)
                except Exception:
                    dist = 1.0
                if dist < 0.0:
                    dist = 0.0
                sim = 1.0 / (1.0 + dist)
                scored.append((sim, doc))

        scored.sort(key=lambda x: x[0], reverse=True)
        top_docs = [doc for _, doc in scored[:TOP_K_FINAL]]

        context = "\n\n---\n\n".join(d.page_content for d in top_docs)
        prompt = f""" 
🎯 Your Goal  
Act as a highly skilled IoT Security Analyst and Network Forensics Expert. Use only the information in the retrieved chunks of multi-source data to answer the user’s question. Do not use general knowledge, external assumptions, or infer answers beyond the retrieved content.  

If the retrieved chunks explicitly contain the answer, provide a clear, concise, and contextually accurate response.  
If the retrieved chunks do not explicitly contain the full answer, respond with the best possible answer based on the available content. 

Respond in plain text without adding any additional instructions, metadata, or commentary. Focus on delivering technically sound, actionable insights where possible.  

{context}

Question: {question}
        """
        response = model.invoke(prompt)

        # Convert AIMessage to string if needed
        pred_text = response.content if hasattr(response, "content") else str(response)

        t1 = time.time()
        mem_after = process.memory_info().rss / (1024**2)
        proc_gpu_after = get_process_gpu_usage()
        device_gpu_after = get_device_gpu_usage()
        cpu_after = process.cpu_times()

        # ✅ CPU% normalized by number of logical CPUs (prevents >100% on multi-core)
        wall = max(t1 - t0, 1e-9)
        cpu_delta = ( (cpu_after.user + cpu_after.system) - (cpu_before.user + cpu_before.system) )
        n_cpus = psutil.cpu_count(logical=True) or 1
        cpu_util_pct = (cpu_delta / wall) * 100.0 / n_cpus

        scores = evaluate_answer(pred_text, answer)
        resp_size, token_count = calculate_response_size(pred_text)

        result_rows.append({
            "question": question,
            "reference": answer,
            "rag_answer": pred_text,
            **{f"rag_{k}": v for k, v in scores.items()}
        })

        metric_rows.append({
            "question": question,
            "elapsed_time": t1 - t0,
            "memory_diff": mem_after - mem_before,
            "cpu_util": cpu_util_pct,  # <-- normalized CPU%
            "proc_gpu_memory_used": proc_gpu_after["gpu_memory_used"] - proc_gpu_before["gpu_memory_used"],
            "device_gpu_memory_used": device_gpu_after["gpu_memory_used"] - device_gpu_before["gpu_memory_used"],
            "device_gpu_utilization": device_gpu_after["gpu_utilization"] - device_gpu_before["gpu_utilization"],
            "response_size": resp_size,
            "token_count": token_count
        })

    df_results = pd.DataFrame(result_rows)
    df_results.to_csv("Evaluation_results/eval_dense_only_llama3.2:3b.csv", index=False)

    df_metrics = pd.DataFrame(metric_rows)
    df_metrics.to_csv("Evaluation_results/system_metrics_dense_only_llama3.2:3b.csv", index=False)

    system_averages = df_metrics.mean(numeric_only=True).to_dict()
    quality_metrics = [col for col in df_results.columns if col.startswith("rag_")]
    quality_averages = df_results[quality_metrics].mean(numeric_only=True).to_dict()

    combined_averages = {f"answer_quality_{k}": v for k, v in quality_averages.items()}
    combined_averages.update({f"system_performance_{k}": v for k, v in system_averages.items()})
    pd.DataFrame([combined_averages]).to_csv("Evaluation_results/combined_averages_dense_only_llama3.2:3b.csv", index=False)

    print("\n📊 Average Answer Quality Metrics:")
    for k, v in quality_averages.items():
        print(f"{k}: {v:.4f}")

    print("\n⚙️ Average System Performance Metrics:")
    for k, v in system_averages.items():
        print(f"{k}: {v:.4f}")

    shutdown_nvml()
    print("\n✅ Dense Only evaluation complete.")

if __name__ == "__main__":
    main()