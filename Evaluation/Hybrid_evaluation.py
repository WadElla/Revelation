# Final transformed script: Hybrid (Full) with both process-specific and device-wide GPU metrics

import os
import time
import psutil
import hashlib
import pandas as pd
from transformers import AutoTokenizer
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from nltk.translate.meteor_score import meteor_score
from rouge_score import rouge_scorer
from bert_score import score as bert_score
from pynvml import (
    nvmlInit, nvmlShutdown, nvmlDeviceGetHandleByIndex,
    nvmlDeviceGetMemoryInfo, nvmlDeviceGetUtilizationRates,
    nvmlDeviceGetComputeRunningProcesses
)
from langchain.schema import Document
from langchain_community.vectorstores import Chroma
from langchain_community.llms import Ollama
from sentence_transformers import CrossEncoder
from rank_bm25 import BM25Okapi
from get_embedding_function import get_embedding_function
from docx import Document as DocxDocument
import torch

from langchain_google_genai import ChatGoogleGenerativeAI
import os
os.environ["GOOGLE_API_KEY"] = "AIzaSyCtowOePpR13zVQ9n2IaKgJSXMBMMNFgG0"

import google.generativeai as genai
genai.configure(api_key="AIzaSyCtowOePpR13zVQ9n2IaKgJSXMBMMNFgG0")


# Config
CHROMA_PATH = None
K_SPARSE = 20
K_DENSE = 30
ALPHA = 0.5
TOP_K_HYBRID = 15
K_KEYWORD = 3
TOP_K_FINAL = 7
MODEL_NAME = "llama3.2:3b"
DOCX_QA_PATH = "Evaluation_results/Backdoor_test3.docx"

# ----------------- Monitoring utils -----------------
def initialize_nvml(): nvmlInit()
def shutdown_nvml(): nvmlShutdown()

def get_process_gpu_usage(gpu_index=0):
    """Get GPU memory usage specific to this Python process."""
    pid = os.getpid()
    handle = nvmlDeviceGetHandleByIndex(gpu_index)
    processes = nvmlDeviceGetComputeRunningProcesses(handle)

    for p in processes:
        if p.pid == pid:
            return {
                "gpu_memory_used": p.usedGpuMemory / (1024 ** 2),  # MB
                "gpu_utilization": None  # Utilization per process is not available
            }
    # If this process isn’t using GPU
    return {"gpu_memory_used": 0.0, "gpu_utilization": None}

def get_device_gpu_usage(gpu_index=0):
    """Get total GPU memory usage and utilization for the entire device."""
    handle = nvmlDeviceGetHandleByIndex(gpu_index)
    mem_info = nvmlDeviceGetMemoryInfo(handle)
    utilization = nvmlDeviceGetUtilizationRates(handle)
    return {
        "gpu_memory_used": mem_info.used / (1024 ** 2),  # MB
        "gpu_utilization": utilization.gpu  # percentage
    }

# ----------------- Retrieval utils ------------------
def keyword_filter_docs(docs, query, top_n=5):
    keywords = query.lower().split()
    scored = [(sum(1 for kw in keywords if kw in doc.page_content.lower()), doc) for doc in docs]
    return [doc for score, doc in sorted(scored, key=lambda x: x[0], reverse=True) if score > 0][:top_n]

def build_bm25_index(docs):
    return BM25Okapi([doc.page_content.lower().split() for doc in docs])

def hybrid_search(query, all_docs, bm25, db):
    # --- Sparse ---
    tokens = query.lower().split()
    sparse_scores = bm25.get_scores(tokens)
    top_sparse = sorted(range(len(all_docs)), key=lambda i: sparse_scores[i], reverse=True)[:K_SPARSE]

    # --- Dense (prefer relevance scores; else distance -> similarity) ---
    scored_dense = []
    try:
        dense_hits = db.similarity_search_with_relevance_scores(query, k=K_DENSE)  # [(doc, rel 0..1)]
        for d, rel in dense_hits:
            try:
                s = float(rel)
            except Exception:
                s = 0.0
            scored_dense.append((d, max(0.0, min(1.0, s))))
    except Exception:
        dense_hits = db.similarity_search_with_score(query, k=K_DENSE)  # [(doc, distance)]
        for d, dist in dense_hits:
            try:
                dist = float(dist)
            except Exception:
                dist = 1.0
            if dist < 0.0:
                dist = 0.0
            sim = 1.0 / (1.0 + dist)  # (0,1], higher = better
            scored_dense.append((d, sim))

    # Build a stable mapping from doc -> index using id or md5(content)
    key_to_idx = {}
    for i, d in enumerate(all_docs):
        did = d.metadata.get("id") if isinstance(d.metadata, dict) and d.metadata.get("id") else None
        key = did if did else hashlib.md5(d.page_content.encode()).hexdigest()
        key_to_idx[key] = i

    merged = {idx: {"bm25": float(sparse_scores[idx]), "vector": 0.0} for idx in top_sparse}
    for d, sim in scored_dense:
        did = d.metadata.get("id") if isinstance(d.metadata, dict) and d.metadata.get("id") else None
        key = did if did else hashlib.md5(d.page_content.encode()).hexdigest()
        idx = key_to_idx.get(key)
        if idx is None:
            continue
        merged.setdefault(idx, {"bm25": 0.0, "vector": 0.0})
        merged[idx]["vector"] = max(merged[idx]["vector"], float(sim))

    # --- Min-max normalize both channels before blending ---
    idxs = list(merged.keys())
    bm25_vals = [merged[i]["bm25"] for i in idxs]
    vec_vals  = [merged[i]["vector"] for i in idxs]

    def _minmax(vals):
        lo, hi = min(vals), max(vals)
        if hi <= lo:
            return [0.5 for _ in vals]
        return [(v - lo) / (hi - lo) for v in vals]

    bm25_n = _minmax(bm25_vals)
    vec_n  = _minmax(vec_vals)

    fused = [(idxs[i], ALPHA * vec_n[i] + (1 - ALPHA) * bm25_n[i]) for i in range(len(idxs))]
    fused.sort(key=lambda x: x[1], reverse=True)

    return [all_docs[idx] for idx, _ in fused]

def deduplicate_docs(docs):
    seen, result = set(), []
    for doc in docs:
        h = hashlib.md5(doc.page_content.encode()).hexdigest()
        if h not in seen:
            seen.add(h)
            result.append(doc)
    return result

def rerank_docs(query, docs, reranker, top_k):
    scores = reranker.predict([(query, d.page_content) for d in docs])
    return [doc for _, doc in sorted(zip(scores, docs), key=lambda t: t[0], reverse=True)][:top_k]
    #[doc for _, doc in sorted(zip(scores, docs), reverse=True)][:top_k]

# ----------------- Answer evaluation -----------------
def evaluate_answer(pred, ref):
    ref_toks, pred_toks = ref.split(), pred.split()
    P, R, F1 = bert_score([pred], [ref], lang="en", verbose=False)
    rouge = rouge_scorer.RougeScorer(['rouge1','rouge2','rougeL'], use_stemmer=True).score(ref, pred)
    return {
        "bert_Precision": P.mean().item(), "bert_Recall": R.mean().item(), "bert_F1": F1.mean().item(),
        "rouge1": rouge['rouge1'].fmeasure, "rouge2": rouge['rouge2'].fmeasure, "rougeL": rouge['rougeL'].fmeasure,
        "bleu": sentence_bleu([ref_toks], pred_toks, smoothing_function=SmoothingFunction().method1),
        "meteor": meteor_score([ref_toks], pred_toks)
    }

def calculate_response_size(text):
    try:
        tokens = AutoTokenizer.from_pretrained("gpt2").tokenize(text)
        return sum(len(t.encode()) for t in tokens), len(tokens)
    except:
        return 0, 0

def load_questions(path):
    doc = DocxDocument(path)
    return {p.text.split(":",1)[0].strip(): p.text.split(":",1)[1].strip()
            for p in doc.paragraphs if ":" in p.text}

# ----------------- Main -----------------
def main():
    initialize_nvml()

    # Load chroma DB path
    global CHROMA_PATH
    if not CHROMA_PATH:
        with open("chroma_sessions/latest_session.txt") as f:
            CHROMA_PATH = f.read().strip()

    embed_fn = get_embedding_function()
    db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embed_fn)
    raw = db.get(include=["documents", "metadatas"])
    all_docs = [Document(page_content=t, metadata=m) for t, m in zip(raw["documents"], raw["metadatas"])]

    bm25 = build_bm25_index(all_docs)
    reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
    
    # Updated model
    #model = ChatGoogleGenerativeAI(model="gemini-2.5-pro", temperature=0.6)
    model = Ollama(model=MODEL_NAME, temperature=0.6)
    

    questions = load_questions(DOCX_QA_PATH)

    result_rows, metric_rows = [], []
    process = psutil.Process()

    for question, answer in questions.items():
        # System monitoring
        mem_before = process.memory_info().rss / (1024**2)
        proc_gpu_before = get_process_gpu_usage()
        device_gpu_before = get_device_gpu_usage()
        cpu_before = process.cpu_times()
        t0 = time.time()

        # Retrieval pipeline
        hybrid = hybrid_search(question, all_docs, bm25, db)[:TOP_K_HYBRID]
        fallback = keyword_filter_docs(all_docs, question, top_n=K_KEYWORD)
        candidates = deduplicate_docs(hybrid + fallback)
        top_docs = rerank_docs(question, candidates, reranker, TOP_K_FINAL)

        # Prompt and inference
        context = "\n\n---\n\n".join(d.page_content for d in top_docs)
        prompt = f""" 
                    🎯 Your Goal  
    Act as a highly skilled IoT Security Analyst and Network Forensics Expert. Use only the information in the retrieved chunks of multi-source data to answer the user’s question. Do not use general knowledge, external assumptions, or infer answers beyond the retrieved content.  

    If the retrieved chunks explicitly contain the answer, provide a clear, concise, and contextually accurate response.  
    If the retrieved chunks do not explicitly contain the full answer, respond with the best possible answer based on the available content. 
  

    Respond in plain text without adding any additional instructions, metadata, or commentary. Focus on delivering technically sound, actionable insights where possible.  

    {context}

    Question: {question}
                """
        response = model.invoke(prompt)

        # ✅ Convert AIMessage to string before evaluation
        pred_text = response.content if hasattr(response, "content") else str(response)

        t1 = time.time()
        mem_after = process.memory_info().rss / (1024**2)
        proc_gpu_after = get_process_gpu_usage()
        device_gpu_after = get_device_gpu_usage()
        cpu_after = process.cpu_times()

        scores = evaluate_answer(pred_text, answer)
        resp_size, token_count = calculate_response_size(pred_text)

        # ✅ CPU% normalized by number of logical CPUs (prevents >100% on multi-core)
        wall = max(t1 - t0, 1e-9)
        cpu_delta = ((cpu_after.user + cpu_after.system) - (cpu_before.user + cpu_before.system))
        n_cpus = psutil.cpu_count(logical=True) or 1
        cpu_util_pct = (cpu_delta / wall) * 100.0 / n_cpus

        result_rows.append({"question": question, "reference": answer, "rag_answer": pred_text, **{f"rag_{k}": v for k,v in scores.items()}})

        metric_rows.append({
            "question": question,
            "elapsed_time": t1 - t0,
            "memory_diff": mem_after - mem_before,
            "cpu_util": cpu_util_pct,
            "proc_gpu_memory_used": proc_gpu_after["gpu_memory_used"] - proc_gpu_before["gpu_memory_used"],
            "device_gpu_memory_used": device_gpu_after["gpu_memory_used"] - device_gpu_before["gpu_memory_used"],
            "device_gpu_utilization": device_gpu_after["gpu_utilization"] - device_gpu_before["gpu_utilization"],
            "response_size": resp_size,
            "token_count": token_count
        })

    # Save per-question results
    df_results = pd.DataFrame(result_rows)
    df_results.to_csv("Evaluation_results/eval_hybrid_full_llama3.2:3b.csv", index=False)

    df_metrics = pd.DataFrame(metric_rows)
    df_metrics.to_csv("Evaluation_results/system_metrics_hybrid_full_llama3.2:3b.csv", index=False)

    # Compute averages
    system_averages = df_metrics.mean(numeric_only=True).to_dict()
    quality_metrics = [col for col in df_results.columns if col.startswith("rag_")]
    quality_averages = df_results[quality_metrics].mean(numeric_only=True).to_dict()

    # Combine and save averages
    combined_averages = {f"answer_quality_{k}": v for k, v in quality_averages.items()}
    combined_averages.update({f"system_performance_{k}": v for k, v in system_averages.items()})
    pd.DataFrame([combined_averages]).to_csv("Evaluation_results/combined_averages_hybrid_full_llama3.2:3b.csv", index=False)

    # Print summary
    print("\n📊 Average Answer Quality Metrics:")
    for k, v in quality_averages.items():
        print(f"{k}: {v:.4f}")

    print("\n⚙️ Average System Performance Metrics:")
    for k, v in system_averages.items():
        print(f"{k}: {v:.4f}")

    shutdown_nvml()
    print("\n✅ Hybrid Full evaluation complete.")

if __name__ == "__main__":
    main()